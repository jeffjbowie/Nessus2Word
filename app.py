import pdb
import glob, os
import csv
from xml.dom import minidom
from docx import Document
from docx.shared import Pt, RGBColor


from operator import itemgetter

''' CVSS3 Vectors '''

cvss3_vectors = {
	'AV' : {
		'N' : 'Network',
		'A' : 'Adjacent Network',
		'L' : 'Local',
		'P' : 'Physical'
	},
	'AC' : {
		'L' : 'Low',
		'H' : 'High',
	},
	'PR' : {
		'N' : 'None',
		'L' : 'Low',
		'H' : 'High',
	},
	'UI' : {
		'N' : 'None',
		'R' : 'Required',
	},
	'S' : {
		'U' : 'Unchanged',
		'C' : 'Changed',
	},
	'C' : {
		'N' : 'None',
		'L' : 'Low',
		'H' : 'High',
	},
	'I' : {
		'N' : 'None',
		'L' : 'Low',
		'H' : 'High',
	},
	'A' : {
		'N' : 'None',
		'L' : 'Low',
		'H' : 'High'
	}
}

vuln_summaries = {
	'Apache 2.4.x < 2.4.42 Multiple Vulnerabilities' : 'The version of Apache found on this host is assocatied with multiple vulnerabilites.',
	'Backup Files Disclosure' : 'Renaming old files on the remote host makes it possible to download and view contents, resulting in sensitive information disclosure.'
}

def save_docx(file_path): 

	# Create a new Word document.
	document = Document()
	font = document.styles['Normal'].font

	# Set font size.
	font.name = 'Arial'
	font.size = Pt(8)

	'''
	Loop through our vulnerabilites, adding paragraphs. 
	Looping through vulnerabilities without calling prepare_vulns() will
	cause separate items for each host , even if multiple share the same vuln.
	It is easier to unique that list, and then add the additional hosts.
	'''

	prepare_vulns()
	sorted_vulns = sorted(unique_vulns, key=itemgetter('severity')) 
	
	for vuln in sorted_vulns:

		# Catch which # of results we are parsing, add 1.clear'
		number = sorted_vulns.index(vuln) + 1

		#Vulnerability Name
		p = document.add_paragraph('')
		p.add_run('%d) %s\n\n' % (number, vuln['name'])).bold = True

		
		p.add_run('CVSS3 Base Score: ')
		_run_base_score = p.add_run(f"{vuln['cvss3_base_score']}\n\n")

		#CVSS3 Info
		vectors = vuln['cvss3_vector'].split('/')[1:]

		attack_vector = cvss3_vectors['AV'][vectors[0].split(':')[1]]
		attack_complexity = cvss3_vectors['AC'][vectors[1].split(':')[1]]
		priveleges_required = cvss3_vectors['PR'][vectors[2].split(':')[1]]
		user_interaction = cvss3_vectors['UI'][vectors[3].split(':')[1]]
		scope = cvss3_vectors['S'][vectors[4].split(':')[1]]
		confidentiality_impact = cvss3_vectors['C'][vectors[5].split(':')[1]]
		integrity_impact = cvss3_vectors['I'][vectors[6].split(':')[1]]
		availability_impact = cvss3_vectors['A'][vectors[7].split(':')[1]]

		# Font size to 6
		font.size = Pt(8)

		p.add_run('Attack Vector: ')
		_run_av = p.add_run(f"{attack_vector}\n")
		p.add_run('Attack Complexity: ')
		_run_ac = p.add_run(f"{attack_complexity}\n")
		p.add_run('Priveleges Required: ')
		_run_pr = p.add_run(f"{priveleges_required}\n")
		p.add_run('User Interaction: ')
		_run_ui = p.add_run(f"{user_interaction}\n")
		p.add_run('Scope: ')
		_run_s = p.add_run(f"{scope}\n")
		p.add_run('Confidentiality Impact: ')
		_run_c = p.add_run(f"{confidentiality_impact}\n")
		p.add_run('Integrity Impact: ')
		_run_i = p.add_run(f"{integrity_impact}\n")
		p.add_run('Availability Impact: ')
		_run_a = p.add_run(f"{availability_impact}\n\n")

		# Back to 8
		font.size = Pt(8)

		# Severity
		p.add_run('Severity: ')
		_run = p.add_run('%s\n' % vuln['severity'].upper())
		_run.font.color.rgb = colors[vuln['severity'].lower()]

		# Probability 
		p.add_run('Probability: ')
		_run2 = p.add_run('%s\n' % vuln['severity'].upper())
		_run2.font.color.rgb = colors[vuln['severity'].lower()]
		
		# Exploit
		p.add_run('Exploit: %s\n\n' % vuln['impact'])
		
		# Hostname /IP 
		p.add_run('Hostname(s)/IP(s): %s\n\n' % vuln['host'])

		# Summary
		if vuln['name'] in vuln_summaries.keys():
			p.add_run('%s \n\n' % vuln_summaries[vuln['name']])
		else:
			p.add_run('%s \n\n' % vuln['summary'])

		# Remediation Steps
		p.add_run('Remediation Steps: \n\n').bold = True	
		
		for i,r in enumerate(vuln['remediations']):
			p.add_run(f"{i+1}. {r.split(':')[1]}\n")


	'''
		CVSS 3.0 Base Score and Recommendations Table
		
		Create a table with headers : 
		Priority | CVSS3.0 Base Score | Observation | Recommendation

	'''

	
	# Sort vulns by CVSS3 Base Score Descending order.
	cvss_sorted_vulns  = sorted(unique_vulns, key=itemgetter('cvss3_base_score'), reverse=True) 

	document.add_page_break()
	
	table = document.add_table(rows = len(cvss_sorted_vulns), cols = 4)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Priority'
	hdr_cells[1].text = 'CVSS3.0 Base Score'
	hdr_cells[2].text = 'Observation'
	hdr_cells[3].text = 'Recommendation'


	for num, vuln in enumerate(cvss_sorted_vulns):
		row_cells = table.add_row().cells
		row_cells[0].text = str(num + 1)
		row_cells[1].text = vuln['cvss3_base_score']
		row_cells[2].text = vuln['name']
		row_cells[3].text = vuln['remediations']
	

	# Save Word document with original file name.
	document.save( file_path.split('\\')[-1:][0].split('.')[0] + '.docx')

def prepare_vulns():
	
	for vuln in vulnerabilities:
		if vuln['name'] not in [ x['name'] for x in unique_vulns ]:
			print(f"{vuln}\n")
			unique_vulns.append(vuln)


def process_nessus(file_path):

	# Grab our entire XML document.
	doc = minidom.parse(file_path)

	# Grab all hosts scanned.
	report_hosts = doc.getElementsByTagName('ReportHost')

	# Loop through list of hosts.
	for host in report_hosts:
	
		# Dictionary containing host details.
		host_details = {'fqdn' : '','ip' :''}
		remediations = []

		# Loop through Host Properties tags and grab IP / FQDN.
		tags = host.getElementsByTagName('HostProperties')[0].getElementsByTagName('tag')

		for tag in tags: 
			if tag.getAttribute('name') == r'host-fqdn':
				host_details.update({'fqdn' : tag.firstChild.data})
			if tag.getAttribute('name') == r'host-ip':
				host_details.update({'ip' : tag.firstChild.data})
			if r'patch-summary-txt' in tag.getAttribute('name') :
				remediations.append(tag.firstChild.data)

		# Grab list of report items.
		report_items = host.getElementsByTagName('ReportItem')

		# Loop through report items.
		for item in report_items:

			# Skip any items that have a Severity of 0 .
			if int(item.getAttribute('severity')) in (0,1,2):
				#print(f"Skipping {item.getElementsByTagName('plugin_name')[0].firstChild.data} due to low severity of {int(item.getAttribute('severity'))}\r\n")
				continue

			if len(item.getElementsByTagName('cvss3_base_score')) == 0:
				continue

			# Iterate results , append to array.
			vulnerabilities.append({
				'name' : item.getElementsByTagName('plugin_name')[0].firstChild.data,
				'severity' : item.getElementsByTagName('risk_factor')[0].firstChild.data.upper(),
				'host' : host_details['ip'],
				'summary' : item.getElementsByTagName('description')[0].firstChild.data.replace('\n',''),
				'impact' : item.getElementsByTagName('synopsis')[0].firstChild.data,
				'remediations' : remediations,
				'cvss3_base_score' : item.getElementsByTagName('cvss3_base_score')[0].firstChild.data,
				'cvss3_vector' : item.getElementsByTagName('cvss3_vector')[0].firstChild.data
			})
	
	prepare_vulns()
	save_docx(file_path)

# List of severity ratings and what colors 
colors = {
	'low' : RGBColor(51, 255, 51),
	'medium' : RGBColor(255,128,0),
	'high' : RGBColor(255,51,51),
	'critical' : RGBColor(186,85,211)
}

# Define working directory, later move to argparse,etc.
WORKING_DIRECTORY = r'C:\Users\User\Code'

# Change into directory
os.chdir(WORKING_DIRECTORY) 

# Scan directory for Nessus results.
for file in os.listdir(WORKING_DIRECTORY):

	# Define blank array to hold our vulnerabilities, clears with each file load.
	vulnerabilities = []
	unique_vulns = []
	
	if file.endswith(".nessus"):
		process_nessus(os.path.join(WORKING_DIRECTORY, file))
